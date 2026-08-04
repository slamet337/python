"""
fix_template.py
===============
Script untuk memperbaiki tag loop pada template .docx agar
tabel menggunakan sintaks {%tr for %} yang benar di docxtpl.

Masalah umum:
  - Template menggunakan {% for item in item_materi %} di dalam sel tabel
  - Hasilnya semua materi muncul dalam SATU sel (terhambur/berantakan)

Solusi:
  - Tag for/endfor harus diletakkan di level ROW tabel
  - Gunakan {%tr for item in item_materi %} di awal baris
  - Gunakan {%tr endfor %} di akhir baris

Cara pakai:
  1. Tutup semua file template di Word terlebih dahulu
  2. Jalankan: python fix_template.py
"""

import os
import re
import shutil
import sys
from pathlib import Path
from docx import Document

# Fix encoding untuk Windows terminal
if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ======================================================
# KONFIGURASI
# ======================================================

BASE_DIR = Path(__file__).parent.resolve()
TEMPLATES_DIR = BASE_DIR / "templates"

# Daftar template yang akan difix (abaikan file .bak)
TEMPLATE_FILES = (
    [p for p in TEMPLATES_DIR.glob("pkkm*.docx") if ".bak" not in p.name]
    + ([BASE_DIR / "pkkm.docx"] if (BASE_DIR / "pkkm.docx").exists() else [])
)


def get_cell_text(cell):
    """Ambil teks dari cell termasuk semua paragraph."""
    return "\n".join(p.text for p in cell.paragraphs)


def replace_tag_in_cell(cell, old_pattern, new_text):
    """Ganti tag di dalam cell di semua paragraph & run."""
    for para in cell.paragraphs:
        # Gabungkan semua run terlebih dahulu untuk menangani split run
        full_text = "".join(run.text for run in para.runs)
        if re.search(old_pattern, full_text):
            new_full = re.sub(old_pattern, new_text, full_text)
            if para.runs:
                para.runs[0].text = new_full
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(new_full)
    return cell


def fix_table_loops(doc_path: Path) -> bool:
    """
    Perbaiki tag for/endfor di tabel template.
    Simpan ke file sementara dulu, lalu replace.
    Return True jika ada perubahan.
    """
    doc = Document(str(doc_path))
    changed = False

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(get_cell_text(c) for c in row.cells)

            # Cek apakah sudah pakai {%tr %} (sudah benar, skip)
            already_tr = bool(re.search(r"\{%\s*tr\s+(for|endfor)", row_text))
            if already_tr:
                print(f"  [OK]  Sudah {{%tr %}} → skip")
                continue

            # Cek for loop inline
            has_for = bool(re.search(r"\{%-?\s*for\s+\w+\s+in\s+", row_text))
            # Cek endfor inline
            has_endfor = bool(re.search(r"\{%-?\s*endfor\s*-?%\}", row_text))

            if has_for:
                for cell in row.cells:
                    cell_text = get_cell_text(cell)
                    if re.search(r"\{%-?\s*for\s+\w+\s+in\s+", cell_text):
                        replace_tag_in_cell(
                            cell,
                            r"\{%-?\s*for\s+(\w+)\s+in\s+(\w+)\s*-?%\}",
                            r"{%tr for \1 in \2 %}"
                        )
                        preview = cell_text[:50].replace("\n", " ").strip()
                        print(f"  [FIX] for  → {{%tr for %}} | '{preview}'")
                        changed = True

            if has_endfor:
                for cell in row.cells:
                    cell_text = get_cell_text(cell)
                    if re.search(r"\{%-?\s*endfor\s*-?%\}", cell_text):
                        replace_tag_in_cell(
                            cell,
                            r"\{%-?\s*endfor\s*-?%\}",
                            r"{%tr endfor %}"
                        )
                        print(f"  [FIX] endfor → {{%tr endfor %}}")
                        changed = True

    if changed:
        # Simpan ke file sementara dulu, baru ganti file asli
        temp_path = doc_path.with_suffix(".tmp.docx")
        backup_path = doc_path.with_name(doc_path.stem + "_backup.docx")

        doc.save(str(temp_path))

        # Backup file asli
        shutil.copy2(str(doc_path), str(backup_path))
        print(f"  [BAK] Backup: {backup_path.name}")

        # Ganti file asli dengan file yang sudah diperbaiki
        os.replace(str(temp_path), str(doc_path))
        print(f"  [OK]  Disimpan: {doc_path.name}")
    else:
        print(f"  [--]  Tidak ada perubahan.")

    return changed


def main():
    print("=" * 60)
    print(" FIX TEMPLATE DOCXTPL - TABLE LOOP TAG")
    print("=" * 60)
    print()

    found_templates = [p for p in TEMPLATE_FILES if p.exists()]

    if not found_templates:
        print("ERROR: Tidak ada template .docx yang ditemukan!")
        print(f"  Dicari di: {TEMPLATES_DIR}")
        print(f"  Dan di   : {BASE_DIR}")
        return

    for tmpl in found_templates:
        print(f"\n>> Memproses: {tmpl.name}")
        try:
            fix_table_loops(tmpl)
        except PermissionError:
            print(f"  [ERR] PERMISSION DENIED!")
            print(f"        Tutup file '{tmpl.name}' di Word, lalu coba lagi.")
        except Exception as e:
            print(f"  [ERR] {e}")

    print()
    print("=" * 60)
    print(" SELESAI!")
    print("=" * 60)
    print()
    print("Struktur tag yang benar di template Word:")
    print()
    print("  Di baris ROW tabel:")
    print("  +--------------------------------------------------------+")
    print("  | {%tr for item in item_materi %}                        |")
    print("  | {{ item.no }} | {{ item.kegiatan }} | {{ item.jp }}    |")
    print("  | {%tr endfor %}                                         |")
    print("  +--------------------------------------------------------+")
    print()
    print("Pastikan juga Total JP di baris BAWAH tabel:")
    print("  {{ total_jp }}")


if __name__ == "__main__":
    main()
