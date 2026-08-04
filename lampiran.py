import os
import re
import shutil
import subprocess
import sys
from pathlib import Path
import docx
import pandas as pd
from docxtpl import DocxTemplate

# =========================================================
# HELPER: HITUNG DURATION / JP DARI WAKTU OTOMATIS
# =========================================================


def hitung_jp_dari_waktu(waktu_str):
    """Menghitung durasi JP secara otomatis dari string 'HH.MM-HH.MM'."""
    if not waktu_str or "-" not in waktu_str:
        return 1

    try:
        match = re.findall(r"(\d{1,2})[\.:](\d{2})", waktu_str)
        if len(match) >= 2:
            jam_mulai, menit_mulai = int(match[0][0]), int(match[0][1])
            jam_selesai, menit_selesai = int(match[1][0]), int(match[1][1])

            total_menit = (jam_selesai * 60 + menit_selesai) - (
                jam_mulai * 60 + menit_mulai
            )

            jp = round(total_menit / 60)
            return max(1, jp)
    except Exception:
        pass

    return 1


def bersihkan_nama(nama):
    if not nama or pd.isna(nama):
        return ""
    nama_str = str(nama)
    gelar_pattern = r"\b(dr|drh|prof|ir|drs|dra|h|hj|s\.h|m\.h|s\.si|m\.t|m\.si|s\.pd|m\.pd|s\.kom|m\.kom|s\.e|m\.m|ll\.m|ph\.d)\b"
    clean = re.sub(gelar_pattern, "", nama_str, flags=re.IGNORECASE)
    clean = re.sub(r"[^a-zA-Z\s]", " ", clean)
    return re.sub(r"\s+", " ", clean).strip().lower()


def is_nama_cocok(nama_excel, nama_jadwal):
    n1 = bersihkan_nama(nama_excel)
    n2 = bersihkan_nama(nama_jadwal)
    if not n1 or not n2:
        return False
    if n1 in n2 or n2 in n1:
        return True
    words1 = set(w for w in n1.split() if len(w) > 2)
    words2 = set(w for w in n2.split() if len(w) > 2)
    if words1 and words2:
        overlap = words1.intersection(words2)
        if len(overlap) >= min(len(words1), len(words2)):
            return True
    return False


# =========================================================
# 1. BACA TEKS DARI FILE DOCX JADWAL
# =========================================================


def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(re.sub(r"\s+", " ", para.text.strip()))

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                re.sub(r"\s+", " ", cell.text.strip().replace("\n", " "))
                for cell in row.cells
            )
            if row_text.strip():
                full_text.append(row_text)

    return "\n".join(full_text)


# =========================================================
# 2. PARSE SEMUA FILE JADWAL DARI FOLDER 'jadwal/'
# =========================================================


def parse_semua_jadwal_folder(jadwal_folder):
    semua_jadwal = []

    for file_path in Path(jadwal_folder).glob("*"):
        if file_path.suffix.lower() in [".docx", ".txt"]:
            print(f" Membaca file jadwal: {file_path.name}")

            if file_path.suffix.lower() == ".docx":
                raw_text = extract_text_from_docx(file_path)
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    raw_text = f.read()

            nama_fakultas = (
                file_path.stem.lower().replace("jadwal", "").strip()
            )
            current_gugus = f"Gugus_{nama_fakultas.upper()}"
            current_kode = ""
            current_ruangan = ""
            current_tanggal = ""

            no_counter = 1
            for line in raw_text.splitlines():
                line = line.strip()
                if not line:
                    continue

                kode_match = re.search(r"\((G\d+K\d+)\)", line)
                if kode_match:
                    current_kode = kode_match.group(1).strip()

                if line.startswith("BT ") or line.startswith("BT"):
                    parts_bt = line.split()
                    if len(parts_bt) >= 2:
                        current_ruangan = f"{parts_bt[0]} {parts_bt[1]}".strip()

                if line.startswith("Hari "):
                    current_tanggal = line.strip()
                    no_counter = 1

                if "|" in line:
                    parts = [p.strip() for p in line.split("|")]
                    if len(parts) >= 4 and not parts[0].startswith("NO"):
                        waktu = parts[1].strip() if len(parts) > 1 else ""
                        kegiatan = parts[2].strip() if len(parts) > 2 else ""
                        pemateri = parts[3].strip() if len(parts) > 3 else ""

                        ruangan_val = (
                            parts[5].strip()
                            if len(parts) > 5 and parts[5].strip()
                            else (
                                parts[4].strip()
                                if len(parts) > 4
                                and not parts[4].replace("-", "").isdigit()
                                else current_ruangan
                            )
                        )

                        no_val = (
                            parts[0]
                            if parts[0].isdigit()
                            else str(no_counter)
                        )

                        if pemateri and "panitia" not in pemateri.lower():
                            jp_hitung = hitung_jp_dari_waktu(waktu)

                            semua_jadwal.append({
                                "fakultas": nama_fakultas,
                                "gugus": current_gugus,
                                "kode": current_kode,
                                "kelas": current_ruangan,
                                "tanggal": current_tanggal,
                                "no": no_val,
                                "waktu": waktu,
                                "kegiatan": kegiatan,
                                "pemateri": pemateri,
                                "ruangan": ruangan_val,
                                "jp": jp_hitung,
                            })
                            no_counter += 1

    return semua_jadwal


def cari_jadwal_pemateri(nama_excel, list_jadwal):
    for item in list_jadwal:
        if is_nama_cocok(nama_excel, item["pemateri"]):
            return item
    return None


def find_libreoffice():
    possible_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for path in possible_paths:
        if Path(path).exists():
            return path
    for command in ["soffice.exe", "soffice", "libreoffice"]:
        found = shutil.which(command)
        if found:
            return found
    return None


def convert_docx_to_pdf(docx_path, output_dir):
    docx_path = Path(docx_path)
    output_dir = Path(output_dir)
    libreoffice = find_libreoffice()

    if not libreoffice:
        print("ERROR: LibreOffice tidak ditemukan!")
        sys.exit(1)

    temp_profile_dir = output_dir / ".libreoffice_tmp"
    profile_url = temp_profile_dir.resolve().as_uri()

    cmd = [
        libreoffice,
        f"-env:UserInstallation={profile_url}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir.resolve()),
        str(docx_path.resolve()),
    ]

    try:
        subprocess.run(
            cmd, capture_output=True, text=True, timeout=180, check=False
        )
    except Exception as e:
        print(f"Gagal Konversi: {e}")
        sys.exit(1)

    return output_dir / f"{docx_path.stem}.pdf"


# =========================================================
# 3. EKSEKUSI UTAMA
# =========================================================


def proses_mail_merge_master():
    base_dir = Path(__file__).parent.resolve()

    excel_path = base_dir / "datapkkm.xlsx"
    folder_jadwal = base_dir / "jadwal"
    folder_templates = base_dir / "templates"
    output_folder = base_dir / "output_pdf"

    output_folder.mkdir(parents=True, exist_ok=True)

    print("=" * 65)
    print(" 1. PROSES BACA SEMUA FILE JADWAL FAKULTAS...")
    print("=" * 65)
    semua_jadwal = parse_semua_jadwal_folder(folder_jadwal)
    print(
        f" Total {len(semua_jadwal)} sesi materi berhasil dimuat dari seluruh jadwal.\n"
    )

    if not excel_path.exists():
        print(f"ERROR: File {excel_path.name} tidak ditemukan!")
        return

    df = pd.read_excel(excel_path)
    df_peserta = df[df["nama"].notnull()].copy()

    print("=" * 65)
    print(f" 2. MEMPROSES {len(df_peserta)} NAMA DOSEN DARI FILE EXCEL MASTER...")
    print("=" * 65)

    for idx, row in enumerate(df_peserta.itertuples(), 1):
        nama = str(getattr(row, "nama", "")).strip()
        mhs = str(getattr(row, "nim", "")) if hasattr(row, "nim") else ""

        # Match nama dosen
        data_jadwal = cari_jadwal_pemateri(nama, semua_jadwal)

        # Hitung Nilai JP
        jp_dosen = data_jadwal["jp"] if data_jadwal else 1
        str_jp = f"{jp_dosen} JP"

        # Pilih Template Word secara Otomatis
        fakultas_key = data_jadwal["fakultas"] if data_jadwal else "default"
        template_spesifik = (
            folder_templates / f"pkkm_{fakultas_key.lower()}.docx"
        )
        template_default = folder_templates / "pkkm.docx"

        if template_spesifik.exists():
            template_path = template_spesifik
        elif template_default.exists():
            template_path = template_default
        else:
            template_path = base_dir / "pkkm.docx"

        # Context Variabel
        context = {
            "nama": nama,
            "mhs": mhs,
            "gugus": data_jadwal["gugus"] if data_jadwal else "",
            "kode": data_jadwal["kode"] if data_jadwal else "",
            "kelas": data_jadwal["kelas"] if data_jadwal else "",
            "tanggal": data_jadwal["tanggal"] if data_jadwal else "",
            "no": "1",  # Selalu mulai dari angka 1
            "waktu": data_jadwal["waktu"] if data_jadwal else "",
            "kegiatan": data_jadwal["kegiatan"] if data_jadwal else "",
            "ruangan": data_jadwal["ruangan"] if data_jadwal else "",
            "jp": str_jp,
            "total_jp": str_jp,
        }

        # Render Word
        doc = DocxTemplate(template_path)
        doc.render(context)

        # --- PENAMAAAN FILE: GUGUS_NAMA ---
        clean_nama = "".join(
            c for c in nama if c.isalnum() or c in (" ", "_", "-")
        ).strip()

        # Gunakan kode gugus (misal G4K1) jika ada, atau nama gugus fakultas
        val_gugus = (
            data_jadwal["kode"]
            if (data_jadwal and data_jadwal["kode"])
            else (data_jadwal["gugus"] if data_jadwal else "Gugus")
        )
        clean_gugus = "".join(
            c for c in val_gugus if c.isalnum() or c in (" ", "_", "-")
        ).strip()

        # Format Nama File: Sertifikat_[Gugus]_[Nama].docx / .pdf
        filename = f"Sertifikat_{clean_gugus}_{clean_nama}.docx"
        docx_out = output_folder / filename
        doc.save(docx_out)

        # Konversi Ke PDF
        pdf_out = convert_docx_to_pdf(docx_out, output_folder)

        if pdf_out.exists():
            docx_out.unlink()  # Hapus temporary docx

        if data_jadwal:
            print(
                f"[{idx}/{len(df_peserta)}] {pdf_out.name} --> (Fakultas: {data_jadwal['fakultas'].upper()}) [PDF OK]"
            )
        else:
            print(
                f"[{idx}/{len(df_peserta)}] ⚠️  {pdf_out.name} --> [JADWAL TIDAK DITEMUKAN!]"
            )

    print("\n" + "=" * 65)
    print(" PROSES SELESAI! Seluruh sertifikat PDF berhasil dibuat.")
    print("=" * 65)


if __name__ == "__main__":
    proses_mail_merge_master()